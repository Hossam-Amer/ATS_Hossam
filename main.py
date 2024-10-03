import os
from docx import Document
import google.generativeai as genai
import PyPDF2 as pdf
import sys
import streamlit as st
import pandas as pd


import ast

# Force output to use utf-8 encoding
sys.stdout.reconfigure(encoding='utf-8')
################################
###important
api_key="PUT YOUR API KEY"

#############################
genai.configure(api_key='api_key')

generation_config = {
    "temperature": 0.1,
    "top_p": 1,
    "top_k": 1,

  }
model=genai.GenerativeModel('gemini-1.0-pro-latest',generation_config=generation_config)

def get_gemini_repsonse(input):
    
  
    
    
    
    response=model.generate_content(input)
    return response.text





def readPdf(uploaded_file):
    reader=pdf.PdfReader(uploaded_file)
    text=""
    for page in range(len(reader.pages)):
        page=reader.pages[page]
        text+=str(page.extract_text())
    return text


def read_docx(file_path):
    """Extract text from DOCX."""
    try:
        doc = Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return None




def compare_cvs_with_job_description(cv_dir, job_desc_text,job_desc_file_path):

    """Compare CVs with job description."""
    cv_text=''
    
    dict={"name": "",  # Full name of the candidate
 "Job Description Match Percentage": "%",  # Percentage value (as a string)
 "Experience Level": "",  # Entry-level, Mid-level, Senior
 "Number of Years of Experience": "",  # Numerical value as string
 "Projects Related only to Job description": "",  # Describe related projects
 "Achievements and Results": "",  # Key achievements relevant to the role   
 "Skills Match with job description": "",  # List matched skills
 "Missing Keywords from the CV": "",  # List missing keywords
 "Keywords Frequency": "",  # Frequency of important keywords
}
    # print(f"Comparing CVs in '{cv_dir}' with job description '{os.path.basename(job_desc_file_path)}':")
    df = pd.DataFrame(columns=list(dict.keys()))

    print(2)

    for cv_file_name in os.listdir(cv_dir):
        cv_file_path = os.path.join(cv_dir, cv_file_name)

        if cv_file_path.endswith('.pdf'):
    
            cv_text = readPdf(cv_file_path)
        elif cv_file_path.endswith('.docx'):
            cv_text = read_docx(cv_file_path)
        
        else:
            print(f"Skipping file {cv_file_name}, unsupported format.")
            
            continue
        input_prompt = f"""
You are a highly skilled ATS (Applicant Tracking System) tasked with evaluating the following resume based on the given job description.
Your goal is to evaluate how well the resume matches the job description exactly and try to be a hard grader.

Return a single Python dictionary, with the following structure:

{{"name": "",  # Full name of the candidate
 "Job Description Match Percentage": "%",  # Percentage value (as a string)
 "Experience Level": "",  # Entry-level, Mid-level, Senior
 "Number of Years of Experience": "",  # Numerical value as string
 "Projects Related only to Job description": "",  # Describe related projects
 "Achievements and Results": "",  # Key achievements relevant to the role   
 "Skills Match with job description": "",  # List matched skills
 "Missing Keywords from the CV": "",  # List missing keywords
 "Keywords Frequency": "",  # Frequency of important keywords
}}

Resume:
{cv_text}

Job Description:
{job_desc_text}

Please ensure the output is returned as a valid Python dictionary where each value is one string.
"""



        print(cv_file_name)
        response=get_gemini_repsonse(input_prompt)
        print(response)
        dict={'name':cv_file_name}
        
        start_index = response.find('{')
        dict_part = response[start_index:]
        dict_part = dict_part.replace('```', '')
        print(dict_part)
        result_dict = ast.literal_eval(dict_part)
        
        
        
        result_dict.update(dict)
        print(df.shape)
        print("___________--------________" )
        print(result_dict )
        print(df )

        df = pd.concat([df, pd.DataFrame([result_dict])], ignore_index=True)
        print(df)
    return df.sort_values(by=['Job Description Match Percentage'],ascending=False,ignore_index=True)
 
        


st.set_page_config(page_title="ATS applcation",layout="wide")
st.header("ATS")
import streamlit as st
import PyPDF2
from io import BytesIO
from docx import Document

def readPdf(file):
    try:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        return f"Error reading PDF: {e}"

def read_docx(file):
    try:
        doc = Document(file)
        text = '\n'.join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error reading DOCX: {e}"

uploaded_file = st.file_uploader("Upload your job description...", type=["pdf", "docx"])

if uploaded_file is not None:
    file_name = uploaded_file.name
    job_desc_text = ''
    msg = ''

    if file_name.endswith(".pdf"):
        job_desc_text = readPdf(BytesIO(uploaded_file.read()))
    elif file_name.endswith(".docx"):
        job_desc_text = read_docx(BytesIO(uploaded_file.read()))

    if job_desc_text:
        # st.write(job_desc_text)
        pass
    else:
        msg = f"Unable to read text from {file_name}."
        st.error(msg)


    print(uploaded_file)

    st.write(msg)
    print(0)
    
    if not msg:
        print(1)
        st.table(
        compare_cvs_with_job_description('CVTest/', job_desc_text,uploaded_file)
            
        )
        
        
